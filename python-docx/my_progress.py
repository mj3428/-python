#!/usr/bin/env python
# encoding: utf-8
'''
@author: miaoj
@contact: major3428@foxmail.com
@software: pycharm
@file: test.py
@time: 2018-10-16 下午 2:54
@desc:
'''

from docx import Document
from docx.shared import Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
import numpy as np
import pandas as pd
from docx.oxml.ns import qn
import datetime


#参数配置
com_name = '恒生印染' #公司名
transformer = '东1#' #变压器名
KVA = 1600000 #变压器容量
startday = '2018年11月23日' #起始时间
endday = '2018年12月4日' #结束时间
freq = 1085 #数据条数
highrisk = 99 #高危数
hiddenrisk = 120 #隐患数
now = datetime.datetime.now() #当前时间的datetime
document = Document('./text/demo.docx')


#南德电气1.82×6.07 电能卫士3.43×9.83
#last_paragraph = document.paragraphs[-1]
#last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 图片居中设置
#开始写入
style_T1 = document.styles.add_style('T1', WD_STYLE_TYPE.PARAGRAPH) #创建T1样式
style_T1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER #居中
document.styles['T1'].font.name = u'宋体' #T1样式使用字体
document.styles['T1']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
document.styles['T1'].font.size = Pt(26)
document.styles['T1'].font.bold = True

style_T2 = document.styles.add_style('T2', WD_STYLE_TYPE.PARAGRAPH) #创建T2样式
style_T2.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER #居中
document.styles['T2'].font.name = u'宋体' #T2样式使用字体
document.styles['T2']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
document.styles['T2'].font.size = Pt(22)
document.styles['T2'].font.bold = True

document.paragraphs[10].style = 'T1'
document.paragraphs[10].add_run(text=com_name+'有限公司')
document.paragraphs[11].clear()
document.paragraphs[11].style = 'T2'
document.paragraphs[11].add_run(text=transformer+'变压器')
document.paragraphs[12].style = 'T2'
document.paragraphs[12].add_run(text='电能健康分析评估报告\n')

#插入印章
#picture = document.add_picture('./pic/poweryun_seal.png', height=Cm(4.00), width=Cm(4.32))
#创建日期样式
style_D1 = document.styles.add_style('Date1', WD_STYLE_TYPE.PARAGRAPH) #创建Date1样式
style_D1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER #居中
document.styles['Date1'].font.name = u'宋体' #T1样式使用字体
document.styles['Date1']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
document.styles['Date1'].font.size = Pt(14)
document.paragraphs[14].clear()
document.paragraphs[14].style = 'Date1'
document.paragraphs[14].add_run(text=now.strftime('%Y')+'年 '+now.strftime('%m')+'月')

#document.add_page_break() #插入分页符
last = document.paragraphs[-1]
last.style = 'Heading 1'
last.add_run(text='一、监测概况及结论')
document.add_heading('1.1电能参数体检结果', level=2)
result = ['—', '↑', '↑', '↑', '—', '↑', '—']
max_value = ['243.9V', '10.1%', '98.6%', '730A', '0.96', '113.5%', '11.9%']
records = (('电压数据', result[0], max_value[0], '205~235V', '《GB/T 12325-2008》'),
          ('谐波电压数据', result[1], max_value[1], '＜5%', '《GB/T14549-1993》'),
          ('电流数据', result[2], max_value[2], '＜In100%', '《JGJ16-2008》'),
          ('谐波电流数据', result[3], max_value[3], '＜各分次国标限值', 'GB/T14549-1993》'),
          ('功率因数', result[4], max_value[4], '0.9~1.0', '《JGJ16-2008》'),
          ('负荷率', result[5], max_value[5], '＜额定容量85%', '《JGJ16-2008》'),
          ('三相电流不平衡度', result[6], max_value[6], '＜15%', '《GB/T 1094-2013》'))
table1 = document.add_table(rows=1, cols=5)
hdr_cells = table1.rows[0].cells
hdr_cells[0].text = '体检项目'
hdr_cells[1].text = '体检结果'
hdr_cells[2].text = '体检值(MAX)'
hdr_cells[3].text = '参考值'
hdr_cells[4].text = '参考标准'
for sty, res, mv, ran, txt in records:
    row_cells = table1.add_row().cells
    row_cells[0].text = sty
    row_cells[1].text = res
    row_cells[2].text = mv
    row_cells[3].text = ran
    row_cells[4].text = txt

widths = [3.99, 2.64, 3.15, 3.53, 4.11]
for i in range(0, 5):
    for cell in table1.columns[i].cells:
        cell.width = Cm(widths[i])

for i in range(0, 8):
    for j in range(0, 5):
        cell = table1.cell(i, j)
        cell_font = cell.paragraphs[0].runs[0].font
        cell_font.color.rgb = RGBColor(0x36, 0x5f, 0x91)
        cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        cell.paragraphs[0].paragraph_format.left_indent = -Cm(0.74)
#cell.font.color.rgb = RGBColor(0x36, 0x5f, 0x91)
table1.style = 'ListCLF'
note = document.add_paragraph("注：（1）‘", style='N1')
fine = note.add_run("—")
fine.bold = True
fine.font.name = 'Arial'
fine.font.color.rgb = RGBColor(0x00, 0xB0, 0x50)
goon = document.paragraphs[-1].add_run("’为标准内，‘")
high = document.paragraphs[-1].add_run("↑")
high.bold = True
high.font.size = Pt(12)
high.font.color.rgb = RGBColor(0xff, 0x00, 0x00)
document.paragraphs[-1].add_run("’为超出标准，‘")
medium = document.paragraphs[-1].add_run("↑")
medium.bold = True
medium.font.size = Pt(12)
medium.font.color.rgb = RGBColor(0xff, 0xff, 0x00)
document.paragraphs[-1].add_run("’为已达临界标准")
document.add_paragraph("（2）以上依据原始导出数据分析，见附件一或用户由平台自行导出", style='N1')
document.add_heading('1.2体检结论说明', level=2)
document.add_paragraph('此结论为%s变压器体检结果，' % transformer +
                        '变压器容量为%dKVA。' % (KVA / 1000) +
                       '%s到%s共采集了%d次数据，高危共计%d次，隐患共120次。' % (startday, endday, highrisk, hiddenrisk) +
                       '以各体检项目在该时段内发生次数及超出各项标准次数进行分析得出：', style='Normal')
document.add_paragraph().add_run('电压体检结论：').bold = True
document.paragraphs[-1].add_run('该变压器在此数据时段内电压值偶尔超出标准，电压数据基本符合要求；' 
                                '谐波电压含量80%以上都超出标准，电网谐波电压含量不合格；')
document.add_paragraph().add_run('电流体检结论：').bold = True
document.paragraphs[-1].add_run('变压器低压侧电流值基本在左右，' 
                                '偶尔会超出变压器额定电流值超载运行，基本属于重载情况；' 
                                '各分次最大谐波电流值都有明显超出标准情况，在均值下5次、7次、11次超出标准，谐波电流不合格；')
document.add_paragraph().add_run('功率因数体检结论：').bold = True
document.paragraphs[-1].add_run('功率因数在0.9~1之间的占比为XX%，功率因数保持良好，功率因数合格；')
document.add_paragraph().add_run('负荷率体检结论：').bold = True
document.paragraphs[-1].add_run('该变压器容量为XXKVA，额定电流约XXA，' 
                                '变压器负荷率基本在XX%左右运行，已超出标准限值，' 
                                '监测期间有发生1级过载（100%~110%）和2级过载（110%~150%），' 
                                '如变压器长期在此状态下工作将会发生安全隐患，变压器负荷率较大；')
document.add_paragraph().add_run('三相电流不平衡体检结论：').bold = True
document.paragraphs[-1].add_run('三相电流不平衡度均在15%以内，概率为XX%，' 
                                '在XX月XX号XX分时发生过一次23%的情况，此时为用电负荷大量下降，' 
                                '属于小电流短暂超标，三相电流不平衡度达标。')
attention = document.add_paragraph(style='N1').add_run('注：分析数据来源为2018年12月1日到2018年12月31日，'
                                 '共采集了2962次（正常情况15分钟为一个周期点，异常时即时推送）')
attention.bold = True
attention.font.size = Pt(7.5)
document.add_page_break() #插入分页符
document.add_heading('1.3体检不正常项目说明', level=2)
document.add_heading('1.3.1谐波电压', level=3)
document.add_paragraph('...')
document.add_heading('1.3.2谐波电流', level=3)
document.add_paragraph('...')
document.add_page_break() #插入分页符
document.add_heading('二、专项体检数据分析说明', level=2)
document.add_paragraph('数据点为%s变压器，变压器容量%dKVA。'% (transformer, KVA / 1000) +
                       '从%s到%s共采集了%d次数据，'%(startday, endday, freq) +
                       '结合各项标准，此时段内存在谐波电压含量、谐波电流含量、负荷率超标隐患。' 
                       '此时段电压数据超出标准限值约XX次；谐波电压含量超出标准限值约XX次；'
                       '电流超出变压器额定电流值约XX次；功率因数超出标准限值约XX次；'
                       '变压器负荷率85%以上约XX次；三相电流不平衡超出标准限值XX次。', style='Normal')
document.add_heading('2.1电压数据', level=2)
document.add_heading('2.1.1 电压数据体检分析', level=3)
document.add_paragraph('根据国家标准《GB/T 12325-2008》中规定单相220V供电电压允许偏差为标称系统电压的+7％、-10%。'
                       '由此计算出电压标准上限值为235.4V。'
                       '从以下分析中可以看出A、B、C三相电压达标率为XX%左右，基本符合标准。',style='Normal')
document.add_heading('1）电压趋势图', level=4)
document.add_paragraph('XXX时段内监测数据，标准限值198V~235V，'
                       '下图中最大值XXX，最小值XXX，用户可在平台内自行查看，一周内每日显示96个数据点，一月内显示每日平均值，'
                       '可根据此图(该图使用2小时均值聚合而成)看出该时段内电压变化状态，用来分析用电情况。', style='Normal')
document.add_picture('./pic/U_trend.png', height=Cm(9.2), width=Cm(17.2))
document.paragraphs[-1].paragraph_format.left_indent = -Cm(0.74) #五号一个字符长度0.371厘米 缩进了两个字符
document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
document.add_paragraph('图2-1  电压趋势图')
document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
document.add_page_break() #插入分页符
document.add_heading('2）电压数据概率分布', level=4)
document.add_picture('./pic/U.png', height=Cm(6.6), width=Cm(17.2)) #电压图17.2*6.6
document.paragraphs[-1].paragraph_format.left_indent = -Cm(0.74)
document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
document.add_paragraph('图2-2 电压值概率分布图')
document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
document.add_heading('3）电压数据合格率', level=4)
document.add_paragraph('取导出数据中各相电压实际超出标准235.4V次数，与总次数对比分析得出该合格率，'
                       '此分析结果可用于查看该时段内各相电压在标准范围的占比，来判断该时段内的电压是否合格。', style='Normal')
document.add_picture('./pic/Upie.png', height=Cm(5.71), width=Cm(17.21)) #环形图17.1*5.71
document.paragraphs[-1].paragraph_format.left_indent = -Cm(0.74)
document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
document.add_paragraph('图2-3 电压合格率环形图')
document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
document.add_paragraph('--------------------------------------------------------------------', style='Normal')
document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.DISTRIBUTE
document.add_heading('2.1.2 电压释义', level=3)
document.add_heading('1）电压的定义', level=4)
document.add_paragraph('电压（voltage），也称作电势差或电位差，是衡量单位电荷在静电场中由于电势不同所产生的能量差的物理量。'
                       '其大小等于单位正电荷因受电场力作用从A点移动到B点所做的功，'
                       '电压的方向规定为从高电位指向低电位的方向。', style='Normal')
document.add_page_break()
document.add_heading('2）电压的算法', level=4)
p = document.add_paragraph('线电压算法:', style='Normal')
r = p.add_run()
r.add_picture('./pic/U_line.png', width=Cm(3.36), height=Cm(1.19))
p = document.add_paragraph('相电压算法:', style='Normal')
r = p.add_run()
r.add_picture('./pic/U_phase.png', width=Cm(1.80), height=Cm(1.19))
document.add_heading('3）参考标准', level=4)
document.add_paragraph('参考《GB/T 12325-2003、2008》电能质量供电电压偏差，'
                       '单相220V供电电压允许偏差为标称系统电压的+7%、-10%。', style='Normal')
document.add_paragraph('--------------------------------------------------------------------', style='Normal')
document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.DISTRIBUTE
document.add_heading('2.2 谐波电压数据', level=2)
document.add_heading('2.2.1 谐波电压体检数据分析小结', level=3)
document.add_paragraph('国家标准《GB/T14549-1993》中相电压最大谐波畸变率限值为5%，从以下分析中可以看出A、B、C各相电压总谐'
                       '波畸变率已明显超出国家5%标准，最大值已达11.5%，电网电压谐波含量需要治理。', style='Normal')
document.add_heading('1）谐波电压趋势图', level=4)
document.add_paragraph('XX月份时段内监测数据，标准限值5%，下图中最大值11.5%，最小值4.2%，用户可在平台内自行查看，'
                       '一周内每日显示96个数据点，一月内显示每日平均值，可根据此图(该图使用2小时均值聚合而成)看出该时段内谐波电压变'
                       '化状态，用来分析谐波电压发生在哪个时段内。', style='Normal')
document.add_picture('./pic/UTHD_trend.png', height=Cm(9.2), width=Cm(17.2))
document.paragraphs[-1].paragraph_format.left_indent = -Cm(0.74)
document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
document.add_paragraph('图2-4  谐波电压含量趋势图')
document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
document.add_heading('2）谐波电压数据详情', level=4)
document.add_paragraph('根据导出数据中各相电压谐波含量发生次数分析得出该分布图，标准限值5%。此分析结果可用于观察该时段内电压谐波'
                       '含量主要分布在多少范围内，并可以看出该范围内谐波电压含量发生的概率，以及超出标准的概率。', style='Normal')
document.add_picture('./pic/U_THD.png', height=Cm(6.6), width=Cm(17.2))
document.paragraphs[-1].paragraph_format.left_indent = -Cm(0.74)
document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
document.add_paragraph('图2-5  谐波电压含量概率分布图')
document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
document.add_heading('3）谐波电压含量合格率', level=4)
document.add_paragraph('取导出数据中各相谐波电压含量实际超出5%标准次数，与总次数对比分析得出该合格率，此分析结果可用于查看该时段内各'
                       '相谐波电压含量在标准范围的占比，来判断该时段内的电压是否合格。', style='Normal')
document.add_picture('./pic/UTHD_pie.png', height=Cm(5.74), width=Cm(17.21))
document.paragraphs[-1].paragraph_format.left_indent = -Cm(0.74)
document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
document.add_paragraph('图2-6 谐波电压含量合格率环形图')
document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
document.add_paragraph('--------------------------------------------------------------------', style='Normal')
document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.DISTRIBUTE
document.add_page_break()
document.add_heading('2.2.2 谐波电压释义', level=3)
document.add_heading('1）谐波电压定义', level=4)
document.add_paragraph('谐波电压是谐波电流流过线路阻抗时产生的，对于特定的配电系统，谐波电流与谐波电压之间的关系如下（欧姆定律）：'
                       '谐波电压=谐波电流*电网阻抗。', style='Normal')
document.add_heading('2）电压总谐波含量算法', level=4)
document.add_picture('./pic/UTHD_percent.png', height=Cm(1.2), width=Cm(3.38))
document.add_paragraph('式中：Uh——第h次谐波电压（方根均值）；\n\t\t  U1——基波电压（方根均值）；', style='Normal')
document.add_heading('3）参考标准', level=4)
document.add_paragraph('依据《GB/T14549-1993》电能质量公用电网谐波国家标准进行。\n\t'
                       '国标《GB/T14549-1993》中公用电网谐波电压（相电压）谐波含量限值：', style='Normal')
table2 = document.add_table(rows=1, cols=4)
hdr_cells2 = table2.rows[0].cells
hdr_cells2[0].text = '电网标称电压\t(kV)'
hdr_cells2[1].text = '电压总谐波畸变率(%)'
hdr_cells2[2].text = '各次谐波电压含有率(%)'
for i in range(0, 3):
    hdr_cells2[i].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    hdr_cells2[i].paragraphs[0].paragraph_format.left_indent = -Cm(0.74)

contents2 = (('奇次', '偶次'),
           ('0.38', '5.0', '4.0', '2.0'))
addr = table2.add_row().cells
addr[2].text = contents2[0][0]
addr[2].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
addr[2].paragraphs[0].paragraph_format.left_indent = -Cm(0.74)
addr[3].text = contents2[0][1]
addr[3].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
addr[3].paragraphs[0].paragraph_format.left_indent = -Cm(0.74)
addr1 = table2.add_row().cells
for j in range(0, 4):
    addr1[j].text = contents2[1][j]
    addr1[j].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    addr1[j].paragraphs[0].paragraph_format.left_indent = -Cm(0.74)
table2widths = [3.07, 3.07, 2.23, 2.23]
for i in range(0, 4):
    for cell in table2.columns[i].cells:
        cell.width = Cm(table2widths[i])

a = table2.cell(0, 0)
b = table2.cell(1, 0)
c = table2.cell(0, 1)
d = table2.cell(1, 1)
e = table2.cell(0, 2)
f = table2.cell(0, 3)
a.merge(b)
c.merge(d)
e.merge(f)
table2.style = 'ListCLF2'
table2.alignment = WD_TABLE_ALIGNMENT.CENTER
document.add_paragraph('表2-1 《GB/T14549-1993》中公用电网谐波电压（相电压）谐波含量限值')
document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
document.add_paragraph('--------------------------------------------------------------------', style='Normal')
document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.DISTRIBUTE
document.add_heading('2.3 电流数据', level=2)
document.add_heading('2.3.1 电流数据体检分析小结', level=3)
document.add_paragraph('该变压器容量为XXXKVA，则二次侧额定电流约为XXXA。监测时间段三相电流基本处于超载状态，最大电流为 '
                       'A相为XXXA，B相为XXXA，C相为XXXA，最高时超出额定电流XXX%左右。', style='Normal')
document.add_page_break()
document.add_heading('1）电流趋势图', level=4)
document.add_paragraph(' 12月份时段内监测数据，标准限值XXA（变压器额定电流），下图中最大值XXA，最小值XXA，用户可在'
                       '平台内自行查看，一周内每日显示96个数据点，一月内显示每日平均值，可根据此图(该图使用2小时均值聚合而成)'
                       '看出该时段内电流变化状态，用来分析用电情况。', style='Normal')
document.add_picture('./pic/I_trend.png', height=Cm(9.2), width=Cm(17.2))
document.paragraphs[-1].paragraph_format.left_indent = -Cm(0.74)
document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
document.add_paragraph('图2-7 电流趋势图')
document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
document.add_heading('2）电流数据详情', level=4)
document.add_paragraph('根据导出数据中各相电流发生次数分析得出该概率图，标准限值454.7A。此分析结果可用于观察该时段内电流主要分'
                       '布在多少范围内，并可以看出该范围内电流值发生的概率，以及超出标准的概率。', style='Normal')
document.add_picture('./pic/I.png', height=Cm(6.6), width=Cm(17.2))
document.paragraphs[-1].paragraph_format.left_indent = -Cm(0.74)
document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
document.add_paragraph('图2-8 电流概率分布图')
document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
document.add_page_break()
document.add_heading('3）电流数据合格率', level=4)
document.add_paragraph('取导出数据中各相电流实际超出454.7A次数，与总次数对比分析得出该合格率，此分析结果可用于查看该时段内各相电'
                       '流在标准范围内、外的占比，来判断该时段内的电流数据是否合格。', style='Normal')
document.add_picture('./pic/I_Pie.png', height=Cm(5.74), width=Cm(17.21))
document.paragraphs[-1].paragraph_format.left_indent = -Cm(0.74)
document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
document.add_paragraph('图2-9 电流合格率环形图')
document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
document.add_paragraph('--------------------------------------------------------------------', style='Normal')
document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.DISTRIBUTE
document.add_heading('2.3.2 电流释义',level=3)
document.add_heading('1）电流定义', level=4)
document.add_paragraph('单位时间里通过导体任一横截面的电量叫做电流强度，简称电流。', style='Normal')
document.add_heading('2）电流算法', level=4)
r0 = document.add_paragraph('电流算法：', style='Normal').add_run()
r0.add_picture('./pic/I_format.png', width=Cm(3.35), height=Cm(1.27))
document.add_heading('3) 参考标准', level=4)
document.add_paragraph('根据国标JGJ16-2008《民用建筑电气设计手册》中4.3.2要求：配变压器长期工作电流不宜大于变压器额'
                       '定电流的85%。', style='Normal')
document.add_paragraph('--------------------------------------------------------------------', style='Normal')
document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.DISTRIBUTE
document.add_heading('2.4 谐波电流数据', level=2)
document.add_heading('2.4.1 谐波电流体检数据分析小结', level=3)
document.add_paragraph('从以下分析中可以看出三相各分次谐波电流值，根据‘国标《GB/T14549-1993》注入公共连接点的谐波电流允许值’中'
                       '规定，该电网谐波电流X次、7X次、X次、X次、X次、XX次、X次、X次均有超出国家标准范围，'
                       '以其中X次、X次谐波为主，X次最大值为216.7A，X次最大值为195.5A，'
                       '电流谐波不达标，需要进行谐波治理。', style='Normal')
document.add_page_break()
document.add_heading('1）谐波电流趋势图', level=4)
document.add_paragraph('12月份时段内监测数据，下图中最大值71.6A，最小值1A，用户可在平台内自行查看，一周内每日显示96个数据点，'
                       '一月内显示每日平均值，可根据此图(该图使用2小时均值聚合而成)看出该时段内谐波电流变化状态，'
                       '用来分析电网谐波含量情况。', style='Normal')
document.add_picture('./pic/ITHD_trend.png',  height=Cm(9.2), width=Cm(17.2))
document.paragraphs[-1].paragraph_format.left_indent = -Cm(0.74)
document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
document.add_paragraph('图2-10 总谐波电流值趋势图')
document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
document.add_heading('2）谐波电流数据详情', level=4)
document.add_paragraph('取导出数据中每相各分次电流谐波最大值与平均值分析得出如下两图。此分析结果可用于观察该时段'
                       '内各分次谐波电流值，主要谐波电流为几次，并可以根据国标《GB/T14549-1993》看出该'
                       '范围内各分次谐波电流值是否超出标准，以判断谐波电流是否合格。', style='Normal')
document.add_picture('./pic/HVmax.png', height=Cm(7.4), width=Cm(16.7))
document.paragraphs[-1].paragraph_format.left_indent = -Cm(0.74)
document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
document.add_paragraph('图2-11 各分次谐波电流最大值')
document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
document.add_page_break()
document.add_picture('./pic/HVmean.png', height=Cm(7.4), width=Cm(16.7))
document.paragraphs[-1].paragraph_format.left_indent = -Cm(0.74)
document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
document.add_paragraph('图2-12 各分次谐波电流平均值')
document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
document.add_heading('各相各分次谐波电流值合格率', level=4)
document.add_paragraph('取导出数据中各相电流分次谐波实际超出标准次数，与总次数对比分析得出该合格率，此分析结果可用于查看该时段'
                       '内各相电流谐波含量在标准范围的占比，来判断该时段内的电流谐波是否合格。', style='Normal')
document.add_picture('./pic/HVqrate.png', height=Cm(7.4), width=Cm(16.7))
document.paragraphs[-1].paragraph_format.left_indent = -Cm(0.74)
document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
document.add_paragraph('图2-13 各分次谐波电流合格率图')
document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
document.add_paragraph('--------------------------------------------------------------------', style='Normal')
document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.DISTRIBUTE
document.add_page_break()
document.add_heading('2.4.2 谐波电流释义', level=3)
document.add_heading('1）谐波电流定义', level=4)
document.add_paragraph('供电系统谐波的定义是对周期性非正弦电量进行傅立叶级数分解，除了得到与电网基波频率相同的分量，还得到一系列'
                       '大于电网基波频率的分量，这部分电量称为谐波。', style='Normal')
document.add_heading('2）电流总谐波含量算法', level=4)
r1 = document.add_paragraph(style='Normal').add_run()
r1.add_picture('./pic/thdi_format.png', width=Cm(3.3), height=Cm(1.2))
document.add_heading('3）参考标准', level=4)
document.add_paragraph('依据《GB/T14549-1993》电能质量公用电网谐波国家标准进行。\n\t'
                       '国标《GB/T14549-1993》中注入公共连接点的谐波电流允许值：',style='Normal')
table3 = document.add_table(rows=1, cols=14)
hdr_cells3 = table3.rows[0].cells
thead3 = ('标准电压（KV）', '基准短路容量（MVA）', '谐波次数及谐波电流允许值（A）')
for i in range(0, 14):
    hdr_cells3[0].text = thead3[0]
    hdr_cells3[1].text = thead3[1]
    hdr_cells3[2].text = thead3[2]
for i in range(0, 3):
    hdr_cells3[i].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    hdr_cells3[i].paragraphs[0].paragraph_format.left_indent = -Cm(0.74)
tr3_1 = ('', '', '3', '5', '7', '9', '11', '13', '15', '17', '19', '21', '23', '25')
tr3_2 = ('0.38', '10', '62', '62', '44', '21', '28', '24', '12', '18', '16', '8.9', '14', '12')
table3_add1 = table3.add_row().cells
for i in range(2, 14):
    table3_add1[i].text = tr3_1[i]
    table3_add1[i].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    table3_add1[i].paragraphs[0].paragraph_format.left_indent = -Cm(0.74)
table3_add2 = table3.add_row().cells
for i in range(0, 14):
    table3_add2[i].text = tr3_2[i]
    table3_add2[i].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    table3_add2[i].paragraphs[0].paragraph_format.left_indent = -Cm(0.74)
g = table3.cell(0, 2)
h = table3.cell(0, 13)
i = table3.cell(0, 0)
j = table3.cell(1, 0)
k = table3.cell(0, 1)
l = table3.cell(1, 1)
g.merge(h)
i.merge(j)
k.merge(l)
for i in range(0, 2):
    for cell in table3.columns[i].cells:
        cell.width = Cm(2.8)
for i in range(2, 14):
    for cell in table3.columns[i].cells:
        cell.width = Cm(1.0)
table3.style = 'ListCLF2'
document.add_paragraph('表2-2 《GB/T14549-1993》中注入公共连接点的谐波电流允许值')
document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
r2 = document.add_paragraph('由于PCC点的短路容量不同于假定基准最小短路容量,应按照国标附录B进行换'
                            '算,换算公式如下：\n\t\t', style='Normal').add_run()
r2.add_picture('./pic/thdi_trans.png', height=Cm(1.25), width=Cm(2.11))
document.add_paragraph('式中∶Sk1:公共连接点的最小短路容量,MVA；\n\t\t  Sk2:基准短路容量,MVA；\n\t\t  Ihp:表2中的第h次谐波'
                       '电流允许值,A；\n\t\t  Ih:短路容量为Sk1时的第h次谐波电流允许值,A。', style='Normal')
document.add_paragraph('--------------------------------------------------------------------', style='Normal')
document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.DISTRIBUTE
document.add_heading('2.5 功率因数数据', level=2)
document.add_heading('2.5.1 总功率因数体检数据分析小结', level=3)
document.add_paragraph('根据国标《JGJ16-2008》中要求，功率因数不宜低于0.9。从以下分析中可以看出功率因数0.9到1的占比为XXX%%，'
                       '0.9以下占比XXX%%，基本符合要求，应到加强对功率因数的巡查监管，避免可能带来的功率因数罚款。', style='Normal')
document.add_page_break()
document.add_heading('1）总功率因数趋势图', level=4)
document.add_paragraph('12月份时段内监测数据，标准限值0.9，下图中最大值XXX，最小值XXX，用户可在平台内自行查看，'
                       '一周内每日显示96个数据点，一月内显示每日平均值，可根据此图(该图使用2小时均值聚合而成)看出该时段内功率因数变'
                       '化状态，用来分析功率因数是否达标。')
document.add_picture('./pic/PF_trend.png', height=Cm(9.1), width=Cm(17.2))
document.paragraphs[-1].paragraph_format.left_indent = -Cm(0.74)
document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
document.add_paragraph('图2-14 总功率因数趋势图')
document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
document.add_heading('2）总功率因数体检数据详情', level=4)
document.add_paragraph('根据导出数据中功率因数发生次数和该时段内超出0.9次数与总次数的对比进行分析。此分析结果可用于观察该时段内功'
                    '率因数主要分布在多少范围内，并可以看出该范围内功率因数发生的概率，以及超出标准的概率，'
                    '可用于分析该时段内功率因数在标准范围内、外的占比，来判断该时段内的功率因数是否合格。', style='Normal')
document.add_picture('./pic/PF.png', height=Cm(6.3), width=Cm(17.2))
document.paragraphs[-1].paragraph_format.left_indent = -Cm(0.74)
document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
document.add_paragraph('图2-15 总功率因数分析图')
document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
document.add_paragraph('--------------------------------------------------------------------', style='Normal')
document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.DISTRIBUTE
document.add_page_break()
document.add_heading('2.5.2功率因数释义', level=3)
document.add_heading('1）功率因数定义', level=4)
document.add_paragraph('所谓功率因数就是有功功率与视在功率的比值，通俗地讲就是用'
                       '电设备的实际出力与用电设备的容量的比值，又简称为力率。', style='Normal')
document.add_heading('2）功率因数的算法', level=4)
document.add_paragraph('在交流电路中，电压与电流之间的相位差(Φ)的余弦叫做功率因数，用符号cosΦ表示，在数值上，功率因数是有功功率和'
                       '视在功率的比值，即cosΦ=P/S，功率因数表是指测量交流电路中有功功率'
                       '与视在功率的比值或测量电压、电流间相位角余弦的电表；', style='Normal')
document.add_picture('./pic/PF_format.png', height=Cm(2.54), width=Cm(8.9))
document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
document.add_heading('3）参考标准', level=4)
document.add_paragraph('（1）国标《JGJ16-2008》民用建筑电气设计规范中第3.6.2条规定：10（6）kV及以下无功补偿宜在配电变压器低压侧'
                       '集中补偿，且功率因数不宜低于0.9。高压侧的功率因数指示应符合当地供电部门的规定。'
                       '\n\t（2）现行的《国家电网公司电力系统电压质量和无功电力管理规定》规定，100kVA及以上10kV供电的电力用户在'
                       '用户在用户高峰负荷时变压器高压侧功率因数不宜低于0.95；其他电力用户功率因数不宜低于0.9。', style='Normal')
document.add_paragraph('--------------------------------------------------------------------', style='Normal')
document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.DISTRIBUTE
document.add_heading('2.6 负荷率数据', level=2)
document.add_heading('2.6.1 负荷率体检数据分析小结', level=3)
document.add_paragraph('从以下分析中可以看出最大负荷率在XXX，负荷率已达XXX左右，属于二级超载，基本是长时间处于超载状态，'
                       '长期在此状态工作可能会发生安全隐患，建议进行整改。', style='Normal')
document.add_heading('1）负荷率趋势图', level=4)
document.add_paragraph('12月份时段内监测数据，标准限值85%，下图中最大值XX，最小值XX%，用户可在平台内自行查看，一周内每日'
                       '显示96个数据点，一月内显示每日平均值，可根据此图(该图使用2小时均值聚合而成)看'
                       '出该时段内负荷率变化状态，用来分析用电情况。', style='Normal')
document.add_page_break()
document.add_picture('./pic/LF_trend.png',  height=Cm(9.1), width=Cm(17.2))
document.paragraphs[-1].paragraph_format.left_indent = -Cm(0.74)
document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
document.add_paragraph('图2-16 变压器负荷率趋势图')
document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
document.add_heading('2）负荷率数据详情', level=4)
document.add_paragraph('根据导出数据中负荷率发生次数和该时段内超出85%次数与总次数的对比进行分析。此分析结果可用于观察该时'
                       '段内变压器负荷率主要分布在多少范围，并可以看出该范围内负荷率发生的概率，以及超出标准的概率，可用于分析该'
                       '时段内负荷率在标准范围内、外的占比，来判断该时段内的负荷率是否合格。', style='Normal')
document.add_picture('./pic/LF.png', height=Cm(6.3), width=Cm(17.0))
document.paragraphs[-1].paragraph_format.left_indent = -Cm(0.74)
document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
document.add_paragraph('图2-17 负荷率数据分析图')
document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
document.add_paragraph('--------------------------------------------------------------------', style='Normal')
document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.DISTRIBUTE
document.add_heading('2.6.2 变压器负荷率释义', level=3)
document.add_heading('1）负荷率定义', level=4)
document.add_paragraph('一定时间内，变压器平均输出的视在功率与变压器额定容量之比称之为负荷率。', style='Normal')
document.add_heading('2）负荷率的算法', level=4)
r3 = document.add_paragraph('负荷率 = S/变压器容量 * 100%；\n\t视在功率', style='Normal').add_run()
r3.add_picture('./pic/S_format.png', height=Cm(0.69), width=Cm(2.76))
document.add_paragraph('总有功功率P = U*I*cosφ\n\t总无功功率Q = U*I*sinφ', style='Normal')
document.add_heading('3）参考标准', level=4)
document.add_paragraph('根据国标JGJ16-2008《民用建筑电气设计手册》中4.3.2要求：配电变压器的长期工'
                       '作负载率不宜大于85%。', style='Normal')
document.add_heading('2.7 三相电流不平衡数据', level=2)
document.add_heading('2.7.1三相电流不平衡体检数据分析小结', level=3)
document.add_paragraph('从以下分析可以看出三相电流不平衡度最大为XX%，结合电流趋势图可以看出此时负载未全部开启，'
                       '属于小电流超标，但在大量负载启用时电流不平衡度也可达XX%，主要是因为C相电流一直高于其他两相，'
                       '导致电流不平衡度被放大，三相电流不平衡度已超出国家标准15%限值，建议对其进行治理。', style='Normal')
document.add_heading('1）三相电流不平衡趋势图', level=4)
document.add_paragraph('XXX时段内监测数据，标准限值15%，下图中最大值145.6%，最小值60.4%，用户可在平台内自行查看，'
                       '一周内每日显示96个数据点，一月内显示每日平均值，可根据此图(该图使用2小时均值聚合而成)'
                       '看出该时段内三相电流不平衡状态，用来分析各相用电情况。', style='Normal')
document.add_picture('./pic/Unb_trend.png', height=Cm(8.93), width=Cm(16.33))
document.paragraphs[-1].paragraph_format.left_indent = -Cm(0.74)
document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
document.add_paragraph('图2-18 三相电流不平衡趋势图')
document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
document.add_page_break()
document.add_heading('3）三相电流不平衡数据详情', level=4)
document.add_paragraph('根据导出数据中三相不平衡度发生次数和该时段内超出15%次数与总次数的对比进行分析。此分析结果可用于观察该'
                       '时段内变压器三相电流平衡度主要分布在多少范围，并可以看出该范围内发生的概率，以及超出标准的概率，'
                       '可用于分析该时段内三相不平衡在标准范围内、外的占比，来判断该时段内的三相电流平衡度是否合格。', style='Normal')
document.add_picture('./pic/Unb.png', height=Cm(6.57), width=Cm(16.9))
document.paragraphs[-1].paragraph_format.left_indent = -Cm(0.74)
document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
document.add_paragraph('图2-19 三相电流不平衡数据分析')
document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
document.add_paragraph('--------------------------------------------------------------------', style='Normal')
document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.DISTRIBUTE
document.add_heading('2.7.2三相不平衡释义', level=3)
document.add_heading('1)变压器三相不平衡定义', level=4)
document.add_paragraph('是指在电力系统中三相电流（或电压）幅值不一致，且幅值差超过规定范围。', style='Normal')
document.add_heading('2)不平衡度的计算公式', level=4)
document.add_picture('./pic/Unb_format.png', height=Cm(1.20), width=Cm(9.63))
document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
document.add_heading('3)参考标准', level=4)
document.add_paragraph('国标《GB/T 1094-2013》下《SD292-1988》第五节 负荷管理 第7.5.2条  变压器的三相负荷力求平衡，不平'
                       '衡度不应大于15%，只带少量单相负荷的三相变压器，零线电流不应超过'
                       '额定电流的25%，不符合上述规定时，应将负荷进行调整。', style='Normal')
document.add_paragraph('--------------------------------------------------------------------', style='Normal')
document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.DISTRIBUTE
document.add_page_break()


document.save('./text/test.docx')

